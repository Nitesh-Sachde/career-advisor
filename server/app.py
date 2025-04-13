import random
from flask import Flask, request, jsonify
from flask import send_file
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from langchain_community.document_loaders import PyPDFLoader
from flask_cors import CORS  # Import CORS
import google.generativeai as genai
import csv
from flask_mail import Mail, Message
import os
from werkzeug.utils import secure_filename
from flask import request, jsonify
from langchain_community.document_loaders import PyPDFLoader
import google.generativeai as genai
import os
import docx
import re
from io import BytesIO
import requests
from urllib.parse import urlencode
import os

CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
# LINKEDIN_REDIRECT_URI = os.getenv("LINKEDIN_REDIRECT_URI")
REDIRECT_URI=os.getenv("REDIRECT_URI")
AUTH_URL=os.getenv("AUTH_URL")
TOKEN_URL=os.getenv("TOKEN_URL")
API_URL=os.getenv("API_URL")

app = Flask(__name__)
CORS(app)
job_description = ""
import re
count = 0
history= {}
question = ""
answer= ""
user_answer=""

data = [
    ["Question", "Actual_Answer", "User_Answer","Score","Selected/Not Selected"]
]
file_name = "interview_analysis.csv"

# Writing to the CSV file
with open(file_name, mode="w", newline="") as file:
    writer = csv.writer(file)
    writer.writerows(data)
print(f"CSV file '{file_name}' created successfully!")



# Define upload folder and ensure it exists
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Store job description globally
job_description = ""

# Google Gemini API Key - consider using environment variables for this in production
GEMINI_API_KEY = "AIzaSyALGEs16JYUEYoeAjxHHKmgSODzVGEO8is"
genai.configure(api_key=GEMINI_API_KEY)


@app.route("/linkedin/login")
def linkedin_login():
    params = {
        "response_type": "code",
        "client_id": CLIENT_ID,
        "redirect_uri": REDIRECT_URI,
        "scope": "openid profile email",
        "state": "secureRandomString123"  # In production, use secure state management!
    }
    auth_url = f"https://www.linkedin.com/oauth/v2/authorization?{urlencode(params)}"
    return jsonify({"auth_url": auth_url})

@app.route('/linkedin/callback')
def callback():
    auth_code = request.args.get('code')
    
    # Exchange the auth code for an access token
    token_data = {
        'grant_type': 'authorization_code',
        'code': auth_code,
        'redirect_uri': REDIRECT_URI,
        'client_id': CLIENT_ID,
        'client_secret': CLIENT_SECRET
    }
    response = requests.post(TOKEN_URL, data=token_data)
    token = response.json().get('access_token')
    
    # Fetch user profile data using the access token
    user_profile = get_linkedin_profile(token)
    
    return jsonify(user_profile)

import requests

API_URL = "https://api.linkedin.com/v2/userinfo"  # Replace with correct LinkedIn API endpoint

def get_linkedin_profile(token):
    headers = {
        "Authorization": f"Bearer {token}",
        "Connection": "Keep-Alive"
    }
    
    try:
        response = requests.get(API_URL, headers=headers)
        
        if response.status_code == 200:
            profile_data = response.json()
            return profile_data
        else:
            return {"error": f"Unable to fetch profile data. Status Code: {response.status_code}, Message: {response.text}"}
    except Exception as e:
        return {"error": f"An error occurred: {str(e)}"}

# Endpoint to fetch network size (for demonstration, use mock data or actual LinkedIn network data)
@app.route('/network/size')
def network_size():
    # You would fetch the network size from LinkedIn's API using the user's access token
    # For now, mock the network size
    network_size_data = {
        "direct_connections": 214,
        "extended_reach": 10799
    }
    return jsonify(network_size_data)

# Endpoint to fetch industry breakdown (mock data for demo)
@app.route('/network/industry')
def industry_breakdown():
    industry_data = [
        {"name": "Technology", "value": 45},
        {"name": "Finance", "value": 20},
        {"name": "Healthcare", "value": 15},
        {"name": "Marketing", "value": 10},
        {"name": "Education", "value": 10}
    ]
    return jsonify(industry_data)

# Endpoint to fetch network growth (mock data for demo)
@app.route('/network/growth')
def network_growth():
    growth_data = {
        "growth_percentage": 18,
        "new_connections": 32,
        "engagement_rate": 5,
        "profile_views": 23
    }
    return jsonify(growth_data)


def extract_email_from_resume(text):
    """Extract email address from resume text using regex"""
    pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
    match = re.search(pattern, text)
    return match.group() if match else None
import networkx as nx
from collections import Counter

def build_network_graph(connections):
    G = nx.Graph()
    for conn in connections:
        G.add_node(conn['name'], industry=conn['industry'], role=conn['role'])
        for other in conn['connections']:
            G.add_edge(conn['name'], other)
    return G

def analyze_roles_industries(graph):
    industries = [graph.nodes[n]['industry'] for n in graph.nodes]
    roles = [graph.nodes[n]['role'] for n in graph.nodes]
    return {
        "industries": dict(Counter(industries)),
        "roles": dict(Counter(roles)),
    }

def recommend_connections(graph):
    recs = []
    for node in graph.nodes:
        neighbors = set(graph.neighbors(node))
        second_degree = set()
        for n in neighbors:
            second_degree.update(graph.neighbors(n))
        suggestions = second_degree - neighbors - {node}
        for s in suggestions:
            recs.append({"from": node, "to": s})
    return recs[:10]

from flask import Blueprint, request, jsonify

network_bp = Blueprint('network', __name__)

@network_bp.route('/api/network/upload', methods=['POST'])
def upload_network():
    data = request.get_json()
    graph = build_network_graph(data['connections'])
    return jsonify({"message": "Graph built successfully", "nodes": list(graph.nodes), "edges": list(graph.edges)})

@network_bp.route('/api/network/analysis', methods=['POST'])
def analyze_network():
    data = request.get_json()
    graph = build_network_graph(data['connections'])

    analysis = {
        "industry_breakdown": analyze_roles_industries(graph)['industries'],
        "role_breakdown": analyze_roles_industries(graph)['roles'],
        "recommendations": recommend_connections(graph),
    }
    return jsonify(analysis)

def extract_text_from_docx(file_path):
    """Extract text from a .docx file using python-docx"""
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def get_resume_text(file_path):
    """Extract text from resume based on file type"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        try:
            loader = PyPDFLoader(file_path)
            pages = loader.load()
            return "".join(page.page_content for page in pages)
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            return ""
            
    elif file_extension in ['.docx', '.doc']:
        try:
            return extract_text_from_docx(file_path)
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            return ""
            
    else:
        # For text files
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {str(e)}")
            return ""

def analyze_resume_for_job(resume_text, job_desc):
    """Use Gemini to analyze how well a resume matches a job description"""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        prompt = f"""
        Analyze how well this resume matches the job description.
        
        Job description:
        {job_desc}
        
        Resume:
        {resume_text}
        
        Provide your analysis in JSON format with the following structure:
        {{
            "match_percentage": a number between 0-100 indicating overall match percentage,
            "strengths": [list of specific strengths the candidate has for this role],
            "weaknesses": [list of key skills or experiences from the job description that the candidate is missing],
            "learning_path": {{
                "title": a specific title for the learning path,
                "description": a brief paragraph describing what the candidate needs to learn,
                "resources": [
                    {{
                        "title": resource title,
                        "type": "course" or "book" or "tutorial" or "project",
                        "link": URL to a real resource,
                        "description": brief description of what this resource will teach
                    }},
                    ... (2-3 resources total)
                ]
            }}
        }}
        
        Include only the JSON in your response, no other text.
        """
        
        response = model.generate_content(prompt)
        
        # Clean the response to get properly formatted JSON
        response_text = response.text
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        try:
            analysis = json.loads(response_text)
            return analysis
        except json.JSONDecodeError:
            print(f"Error parsing JSON: {response_text}")
            return create_fallback_analysis()
            
    except Exception as e:
        print(f"Error analyzing resume for job: {str(e)}")
        return create_fallback_analysis()

def create_fallback_analysis():
    """Create a fallback analysis if Gemini API fails"""
    return {
        "match_percentage": 50,
        "strengths": [
            "Has some relevant technical skills",
            "Previous work experience in the industry"
        ],
        "weaknesses": [
            "Missing specific technology expertise",
            "Limited experience with required tools"
        ],
        "learning_path": {
            "title": "Technical Skills Development Path",
            "description": "Focus on building technical skills most relevant to this position",
            "resources": [
                {
                    "title": "Online Learning Platform Course",
                    "type": "course",
                    "link": "https://www.coursera.org/",
                    "description": "Take courses on the specific technologies mentioned in the job description"
                },
                {
                    "title": "Practice Project",
                    "type": "project",
                    "link": "https://github.com/",
                    "description": "Build a portfolio project showcasing the skills needed for this role"
                }
            ]
        }
    }



def similarity_checker_tfidf_job(job_desc, resumes_folder):
    """Calculate similarity scores between job description and resumes using TF-IDF"""
    similarity_scores = []

    for resume_file in os.listdir(resumes_folder):
        resume_path = os.path.join(resumes_folder, resume_file)
        
        # Skip folders and non-resume files
        if os.path.isdir(resume_path) or not resume_file.lower().endswith(('.pdf', '.docx', '.doc', '.txt')):
            continue
            
        resume_text = get_resume_text(resume_path)
        if not resume_text:
            continue
            
        # Calculate TF-IDF similarity
        try:
            corpus = [job_desc, resume_text]
            vectorizer = TfidfVectorizer()
            tfidf_matrix = vectorizer.fit_transform(corpus)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
            
            # Extract email from resume
            email = extract_email_from_resume(resume_text)
            if not email:
                email = f"candidate_{resume_file.split('.')[0]}@example.com"
                
            # Get detailed analysis from Gemini
            analysis = analyze_resume_for_job(resume_text, job_desc)
            
            # Use the match percentage from Gemini if available, otherwise use TF-IDF similarity
            match_score = analysis.get("match_percentage", similarity[0][0] * 100) / 100
            
            result = {
                "resume": resume_file,
                "similarity_score": match_score,
                "email": email,
                "strengths": analysis.get("strengths", []),
                "weaknesses": analysis.get("weaknesses", []),
                "learning_path": analysis.get("learning_path", {})
            }
            
            similarity_scores.append(result)
        except Exception as e:
            print(f"Error processing {resume_file}: {str(e)}")
    
    # Sort by similarity score in descending order
    similarity_scores.sort(key=lambda x: x["similarity_score"], reverse=True)
    return similarity_scores

@app.route('/job_description_job', methods=['POST'])
def set_job_description():
    """Store the job description for analysis"""
    try:
        data = request.json
        global job_description
        job_description = data.get('job_description', '')
        return jsonify({"message": "Job description received", "job_description": job_description})
    except Exception as e:
        print(f"Error in job_description route: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/check_similarity_job', methods=['GET'])
def check_similarity_job():
    """Compare resumes with the job description and return similarity scores"""
    try:
        global job_description
        if not job_description:
            return jsonify({"error": "No job description set. Please submit a job description first."}), 400
            
        similarity_scores = similarity_checker_tfidf_job(job_description, app.config['UPLOAD_FOLDER'])
        return jsonify(similarity_scores)
    except Exception as e:
        print(f"Error in check_similarity route: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/upload_resumes_job', methods=['POST'])
def upload_resumes_job():
    """Handle resume file uploads"""
    try:
        if 'resumes' not in request.files:
            return jsonify({"error": "No files uploaded."}), 400

        files = request.files.getlist('resumes')
        file_paths = []
        
        for file in files:
            if file and file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                file_paths.append(file_path)

        return jsonify({"message": "Files uploaded successfully", "file_paths": file_paths}), 200
    except Exception as e:
        print(f"Error in upload_resumes route: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/view_resume_job', methods=['GET'])
def view_resume_job():
    """Retrieve a resume file for viewing"""
    try:
        file_name = request.args.get('file')
        if not file_name:
            return jsonify({"error": "No file specified"}), 400
            
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_name)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
            
        return send_file(file_path)
    except Exception as e:
        print(f"Error in view_resume route: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/generate_learning_path_job', methods=['POST'])
def generate_learning_path_job():
    """Generate personalized learning path for a candidate"""
    try:
        data = request.json
        resume_file = data.get('resume')
        email = data.get('email')
        
        if not resume_file or not job_description:
            return jsonify({"error": "Missing resume file or job description"}), 400
            
        resume_path = os.path.join(app.config['UPLOAD_FOLDER'], resume_file)
        if not os.path.exists(resume_path):
            return jsonify({"error": "Resume file not found"}), 404
            
        resume_text = get_resume_text(resume_path)
        analysis = analyze_resume_for_job(resume_text, job_description)
        
        return jsonify(analysis)
    except Exception as e:
        print(f"Error in generate_learning_path route: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Simple health check endpoint"""
    return jsonify({"status": "ok"})



def extract_email_from_resume(text):
    pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
    match = re.search(pattern, text)
    return match.group() if match else None

def similarity_checker_tfidf(job_description, resumes_folder):
    similarity_scores = []

    for resume_file in os.listdir(resumes_folder):
        if resume_file.endswith(".pdf"):
            resume_path = os.path.join(resumes_folder, resume_file)
            loader = PyPDFLoader(resume_path)
            resume_text = "".join(page.page_content for page in loader.load())
            corpus = [job_description, resume_text]
            vectorizer = TfidfVectorizer()
            tfidf_matrix = vectorizer.fit_transform(corpus)
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
            email = extract_email_from_resume(resume_text)
            similarity_scores.append((resume_file, similarity[0][0], email))
    
    return similarity_scores



# Configure the API key for Google Generative AI
genai.configure(api_key="AIzaSyALGEs16JYUEYoeAjxHHKmgSODzVGEO8is")

def extract_text_from_docx(file_bytes):
    """Extract text from a .docx file using python-docx"""
    doc = docx.Document(BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

@app.route('/api/analyze-resume', methods=['POST'])
def analyze_resume():
    if 'resume' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['resume']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    # Process different file types
    try:
        resume_text = ""
        file_extension = os.path.splitext(file.filename)[1].lower()
        
        if file_extension == '.pdf':
            # Save PDF temporarily to extract text with PyPDFLoader
            temp_path = os.path.join(app.config['UPLOAD_FOLDER'], "temp_" + file.filename)
            file.save(temp_path)
            loader = PyPDFLoader(temp_path)
            pages = loader.load()
            resume_text = "".join(page.page_content for page in pages)
            os.remove(temp_path)  # Clean up temporary file
            
        elif file_extension in ['.docx', '.doc']:
            # Process Word documents
            resume_text = extract_text_from_docx(file.read())
            
        else:
            return jsonify({"error": "Unsupported file format. Please upload a PDF or Word document."}), 400
        
        # Use Gemini to analyze the resume
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        prompt = f"""
        Analyze this resume text and provide suggestions for improvement. Focus on:
        1. Action verbs and impactful language
        2. Quantifiable achievements
        3. Technical skills presentation
        4. ATS optimization
        5. Clarity and conciseness
        
        For each suggestion, provide:
        - The original text
        - A suggested improvement
        - The reason for the change
        
        Format your response as valid JSON with the following structure:
        {{
            "text": "the full resume text",
            "suggestions": [
                {{
                    "startIndex": start character index in the text,
                    "endIndex": end character index in the text,
                    "original": "original text",
                    "suggestion": "improved text",
                    "reason": "explanation for the improvement"
                }},
                ...more suggestions...
            ]
        }}
        
        Include 3-7 high-impact suggestions. Only include the JSON in your response, no other text.
        
        Resume text:
        {resume_text}
        """
        
        response = model.generate_content(prompt)
        
        # Clean the response to get properly formatted JSON
        response_text = response.text
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
        
        # Parse and return the analysis results
        import json
        try:
            analysis_results = json.loads(response_text)
            return jsonify(analysis_results)
        except json.JSONDecodeError as e:
            # Create a simplified response with the error
            print(f"JSON parse error: {str(e)}")
            print(f"Raw response: {response_text}")
            
            # Extract suggestions using regex if JSON parsing fails
            fallback_results = {
                "text": resume_text,
                "suggestions": extract_fallback_suggestions(resume_text)
            }
            
            return jsonify(fallback_results)
        
    except Exception as e:
        print(f"Error analyzing resume: {str(e)}")
        return jsonify({"error": f"Resume analysis failed: {str(e)}"}), 500

def extract_fallback_suggestions(resume_text):
    """Generate fallback suggestions if JSON parsing fails"""
    # Look for common resume improvement opportunities
    suggestions = []
    
    # Example: Find passive voice phrases
    passive_patterns = [
        (r"(is|are|was|were|been|be|being) (managed|developed|created|designed|implemented)",
         "Use active voice instead of passive voice for stronger impact")
    ]
    
    # Example: Look for weak action verbs
    weak_verbs = ["worked on", "helped with", "assisted", "responsible for"]
    
    index = 0
    for pattern, reason in passive_patterns:
        for match in re.finditer(pattern, resume_text, re.IGNORECASE):
            start, end = match.span()
            original = match.group(0)
            
            # Create a simple improvement suggestion
            action_verb = original.split()[1].replace('ed', '')
            suggestion = action_verb.capitalize() + "ed"
            
            suggestions.append({
                "startIndex": start,
                "endIndex": end,
                "original": original,
                "suggestion": suggestion,
                "reason": reason
            })
            index += 1
            if index >= 3:  # Limit to 3 fallback suggestions
                break
    
    # Check for weak verbs
    for verb in weak_verbs:
        for match in re.finditer(r"\b" + re.escape(verb) + r"\b", resume_text, re.IGNORECASE):
            start, end = match.span()
            
            # Suggest stronger alternatives
            alternatives = {
                "worked on": "Developed",
                "helped with": "Contributed to",
                "assisted": "Supported",
                "responsible for": "Managed"
            }
            
            suggestions.append({
                "startIndex": start,
                "endIndex": end,
                "original": match.group(0),
                "suggestion": alternatives.get(verb.lower(), "Accomplished"),
                "reason": "Replace weak or passive phrases with strong action verbs"
            })
            index += 1
            if index >= 5:  # Limit total suggestions
                break
    
    # If we still don't have enough suggestions, add a generic one
    if len(suggestions) == 0:
        suggestions.append({
            "startIndex": 0,
            "endIndex": min(50, len(resume_text)),
            "original": resume_text[:min(50, len(resume_text))],
            "suggestion": "Start your resume with a compelling professional summary",
            "reason": "A strong summary helps recruiters quickly understand your value proposition"
        })
    
    return suggestions




def calculate_similarity(actual_answer, user_answer):
    """
    Calculate the similarity between the actual answer and the user's answer using TF-IDF and Cosine Similarity.
    
    Args:
    actual_answer (str): The actual correct answer.
    user_answer (str): The answer provided by the user.
    
    Returns:
    float: Similarity score between 0 and 1.
    """
    if not actual_answer or not user_answer:
        return 0.0  # Return 0 if any of the answers is empty
    
    # Create a TF-IDF Vectorizer
    vectorizer = TfidfVectorizer()
    
    # Transform the answers into TF-IDF vectors
    tfidf_matrix = vectorizer.fit_transform([actual_answer, user_answer])
    
    # Compute Cosine Similarity between the vectors
    similarity_score = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0]
    
    return similarity_score


@app.route('/job_description',methods=['POST'])
def job_description():
    data = request.json
    global job_description
    job_description = data.get('job_description', '')
    print("Gotten "+job_description)
    return jsonify(job_description)
    
@app.route('/check_similarity', methods=['GET'])
def check_similarity():
    similarity_scores = similarity_checker_tfidf(job_description, UPLOAD_FOLDER)
    response_data = [{"resume": file, "similarity_score": score, "email": email} for file, score, email in similarity_scores]
    return jsonify(response_data)





@app.route('/view_resume', methods=['GET'])
def view_resume():
    file_name = request.args.get('file')
    file_path = os.path.join(UPLOAD_FOLDER, file_name)
    return send_file(file_path)



# Define the path where resumes will be saved
UPLOAD_FOLDER = "uploads"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/upload_resumes', methods=['POST'])
def upload_resumes():
    if 'resumes' not in request.files:
        return jsonify({"error": "No files uploaded."}), 400

    files = request.files.getlist('resumes')
    file_paths = []
    for file in files:
        if file and file.filename:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)  # Save the file to the specified path
            file_paths.append(file_path)

    return jsonify({"message": "Files uploaded successfully", "file_paths": file_paths}), 200


# # Configure the API key for Google Generative AI

# # Define a route to generate the MCQs
@app.route('/generate_mcqs', methods=['GET'])
def generate_mcqs():
    genai.configure(api_key="AIzaSyALGEs16JYUEYoeAjxHHKmgSODzVGEO8is")
    try:
        
        difficulty = request.args.get('difficulty', 'medium')
        tech_stack = request.args.get('tech_stack', 'general')
        
        model = genai.GenerativeModel("gemini-1.5-flash")
        print("Difficulty:", difficulty)
        print("Categorization:", tech_stack)

        # Create prompt with both parameters
        prompt = f"Generate 10 {difficulty} level MCQs of {tech_stack} with correct option. Give me output in JSON Format with keys question and options in A,B,C,D and answer."

        response = model.generate_content(prompt)
        print(response.text)
        # Return the generated response as JSON
        return jsonify({"mcqs": response.text})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/extract_tech_stack', methods=['POST'])
def extract_tech_stack():
    try:
        genai.configure(api_key="AIzaSyALGEs16JYUEYoeAjxHHKmgSODzVGEO8is")
        
        # Check if the request contains a file
        if 'resume' not in request.files:
            return jsonify({"error": "No resume file provided"}), 400
            
        file = request.files['resume']
        
        # Check if a file was actually selected
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
            
        # Save the file temporarily to read its contents
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], "temp_" + file.filename)
        file.save(temp_path)
        
        # Extract text from the file based on file type
        if file.filename.endswith('.pdf'):
            loader = PyPDFLoader(temp_path)
            pages = loader.load()
            resume_text = "".join(page.page_content for page in pages)
        elif file.filename.endswith(('.docx', '.doc')):
            # You would need to add docx handling library here
            # For example using python-docx or docx2txt
            return jsonify({"error": "DOCX format currently not supported"}), 400
        else:
            # For plain text files
            with open(temp_path, 'r', encoding='utf-8') as f:
                resume_text = f.read()
        
        # Remove the temporary file
        os.remove(temp_path)
        
        # Log the extracted text
        print("Extracted Resume Text:", resume_text[:200] + "..." if len(resume_text) > 200 else resume_text)
        
        # Use Gemini model to extract tech stack
        model = genai.GenerativeModel("gemini-1.5-flash")
        prompt = f"From the following resume text, extract the technologies, frameworks, programming languages, and tools the candidate has experience with. Return only a JSON array of strings with the names of the technologies: {resume_text}"

        response = model.generate_content(prompt)
        tech_stack_text = response.text
        
        # Clean the response to get a properly formatted JSON
        # Remove any markdown code block syntax if present
        if tech_stack_text.startswith('```json'):
            tech_stack_text = tech_stack_text.replace('```json', '').replace('```', '').strip()
        elif tech_stack_text.startswith('```'):
            tech_stack_text = tech_stack_text.replace('```', '').strip()
            
        try:
            # Attempt to parse the tech stack as JSON
            tech_stack = json.loads(tech_stack_text)
            
            # Determine a likely category based on the tech stack
            categories = {
                "Technology": ["javascript", "python", "java", "c#", "react", "angular", "vue", "node", "express", "django", "flask", "spring", "aws", "azure", "devops", "cloud", "docker", "kubernetes", "database", "sql", "nosql", "mongodb", "postgresql", "mysql"],
                "Design": ["figma", "sketch", "adobe", "photoshop", "illustrator", "ui", "ux", "design", "graphic", "visual", "interface"],
                "Business": ["excel", "powerpoint", "word", "office", "tableau", "powerbi", "analytics", "management", "marketing", "sales", "finance", "accounting"]
            }
            
            category = "Technology"  # Default category
            
            # Simple logic to determine category based on keywords in tech stack
            tech_string = " ".join(tech_stack).lower()
            for cat, keywords in categories.items():
                if any(keyword in tech_string for keyword in keywords):
                    category = cat
                    break
                    
            print("Extracted Tech Stack:", tech_stack)
            
            return jsonify({
                "techStack": tech_stack,
                "category": category
            })
            
        except json.JSONDecodeError as e:
            # If parsing fails, return the raw text
            return jsonify({
                "techStack": tech_stack_text.split(","),
                "category": "Technology"
            })
    
    except Exception as e:
        print(f"Error in extract_tech_stack: {str(e)}")
        return jsonify({"error": str(e)}), 500


import json

@app.route('/generate_learning_path', methods=['POST'])
def generate_learning_path():
    try:
        # Get data from request
        data = request.json
        
        if not data:
            return jsonify({"error": "No data provided"}), 400
            
        # Extract assessment information
        score = data.get('score', 0)
        difficulty = data.get('difficulty', 'beginner')
        tech_stack = data.get('techStack', [])
        question_answers = data.get('questionAnswers', [])
        
        # Analyze incorrect answers to identify knowledge gaps
        knowledge_gaps = analyze_knowledge_gaps(question_answers)
        
        # Generate personalized learning path based on assessment
        learning_path = create_learning_path(score, difficulty, tech_stack, knowledge_gaps)
        
        return jsonify({"learningPath": learning_path})
        
    except Exception as e:
        print(f"Error generating learning path: {str(e)}")
        return jsonify({"error": str(e)}), 500

def analyze_knowledge_gaps(question_answers):
    """
    Analyze question answers to identify knowledge gaps
    """
    knowledge_gaps = []
    
    # Extract topics from incorrect answers
    for qa in question_answers:
        if not qa.get('isCorrect', False):
            # Extract keywords or topics from the question
            question_text = qa.get('question', '').lower()
            
            # Simple keyword extraction (enhance this with NLP in production)
            keywords = extract_keywords_from_question(question_text)
            knowledge_gaps.extend(keywords)
    
    # Count occurrences of each keyword to identify most common gaps
    gap_counts = {}
    for gap in knowledge_gaps:
        if gap in gap_counts:
            gap_counts[gap] += 1
        else:
            gap_counts[gap] = 1
    
    # Sort gaps by frequency
    sorted_gaps = sorted(gap_counts.items(), key=lambda x: x[1], reverse=True)
    
    # Return top knowledge gaps
    return [gap[0] for gap in sorted_gaps[:3]]

def extract_keywords_from_question(question_text):
    """
    Extract relevant keywords from a question
    This is a simple implementation - in production, use NLP techniques
    """
    # Define common technical keywords to look for
    tech_keywords = [
        "javascript", "react", "vue", "angular", "node", "express", 
        "api", "rest", "graphql", "database", "sql", "nosql", "mongodb",
        "python", "django", "flask", "machine learning", "data science",
        "algorithm", "data structure", "oop", "functional programming",
        "frontend", "backend", "full stack", "devops", "docker", "kubernetes",
        "aws", "cloud", "security", "authentication", "testing"
    ]
    
    found_keywords = []
    for keyword in tech_keywords:
        if keyword in question_text:
            found_keywords.append(keyword)
    
    return found_keywords

def create_learning_path(score, difficulty, tech_stack, knowledge_gaps):
    """
    Create a personalized learning path based on assessment results
    """
    # Define difficulty-based learning approach
    difficulty_approaches = {
        "beginner": {
            "title": "Foundational Learning Path",
            "description": "Build a strong foundation in core concepts and practical applications."
        },
        "intermediate": {
            "title": "Skill Enhancement Path",
            "description": "Deepen your understanding and develop specialized expertise."
        },
        "advanced": {
            "title": "Expert Mastery Path",
            "description": "Master advanced concepts and become an industry expert."
        }
    }
    
    # Get base path structure from difficulty
    base_path = difficulty_approaches.get(difficulty.lower(), difficulty_approaches["intermediate"])
    
    # Create title based on tech stack focus
    tech_focus = tech_stack[0] if tech_stack else "Technology"
    base_path["title"] = f"{base_path['title']}: {tech_focus} Specialization"
    
    # Add personalized description based on score
    if score < 40:
        base_path["description"] = f"This learning path is designed to strengthen your foundational knowledge in {tech_focus} and address key knowledge gaps."
    elif score < 70:
        base_path["description"] = f"Build upon your existing {tech_focus} skills with this focused learning path targeting your specific areas for improvement."
    else:
        base_path["description"] = f"Take your {tech_focus} expertise to the next level with advanced topics and specialized knowledge areas."
    
    # Generate resources based on tech stack and knowledge gaps
    resources = generate_resources(tech_stack, knowledge_gaps, difficulty, score)
    
    # Create complete learning path object
    learning_path = {
        "title": base_path["title"],
        "description": base_path["description"],
        "resources": resources
    }
    
    return learning_path

def generate_resources(tech_stack, knowledge_gaps, difficulty, score):
    """
    Generate specific learning resources based on assessment
    """
    # Technology-specific resource database (expand this in production)
    resource_database = {
        "javascript": [
            {
                "title": "JavaScript Fundamentals",
                "type": "course",
                "link": "https://javascript.info/",
                "description": "Master core JavaScript concepts and modern features."
            },
            {
                "title": "Build 30 JavaScript Projects in 30 Days",
                "type": "project",
                "link": "https://javascript30.com/",
                "description": "Apply JavaScript through hands-on projects."
            }
        ],
        "react": [
            {
                "title": "React Complete Guide",
                "type": "course",
                "link": "https://reactjs.org/tutorial/tutorial.html",
                "description": "Learn React from fundamentals to advanced patterns."
            },
            {
                "title": "Build a React Dashboard",
                "type": "project",
                "link": "https://github.com/topics/react-dashboard",
                "description": "Create a data visualization dashboard with React."
            }
        ],
        "python": [
            {
                "title": "Python for Everybody",
                "type": "course",
                "link": "https://www.py4e.com/",
                "description": "Comprehensive Python programming fundamentals."
            },
            {
                "title": "Build a Python Web Scraper",
                "type": "project",
                "link": "https://realpython.com/beautiful-soup-web-scraper-python/",
                "description": "Learn web scraping techniques with Python."
            }
        ],
        "api": [
            {
                "title": "RESTful API Design",
                "type": "tutorial",
                "link": "https://restfulapi.net/",
                "description": "Learn best practices for designing RESTful APIs."
            }
        ],
        "database": [
            {
                "title": "SQL Masterclass",
                "type": "course",
                "link": "https://www.sqlcourse.com/",
                "description": "Learn fundamental and advanced SQL concepts."
            }
        ],
        "machine learning": [
            {
                "title": "Machine Learning Foundations",
                "type": "course",
                "link": "https://www.coursera.org/learn/machine-learning",
                "description": "Build strong foundations in machine learning theory and practice."
            }
        ]
    }
    
    # Resources for common knowledge gaps not tied to specific technologies
    general_resources = [
        {
            "title": "Computer Science Fundamentals",
            "type": "course",
            "link": "https://www.edx.org/course/introduction-computer-science-harvardx-cs50x",
            "description": "Build a strong foundation in computer science concepts."
        },
        {
            "title": "Algorithm Practice",
            "type": "tutorial",
            "link": "https://leetcode.com/",
            "description": "Sharpen your algorithm and problem-solving skills."
        },
        {
            "title": "System Design Fundamentals",
            "type": "course",
            "link": "https://github.com/donnemartin/system-design-primer",
            "description": "Learn how to design scalable systems and applications."
        }
    ]
    
    resources = []
    
    # Add technology-specific resources based on tech stack
    for tech in tech_stack:
        tech_lower = tech.lower()
        if tech_lower in resource_database:
            # Add 1-2 resources per technology
            tech_resources = resource_database[tech_lower]
            selected = tech_resources if len(tech_resources) <= 2 else random.sample(tech_resources, 2)
            resources.extend(selected)
    
    # Add resources for knowledge gaps
    for gap in knowledge_gaps:
        if gap in resource_database:
            gap_resources = resource_database[gap]
            selected = gap_resources[0] if gap_resources else None
            if selected and selected not in resources:
                resources.append(selected)
    
    # If we don't have enough resources, add general ones
    if len(resources) < 3:
        remaining = 3 - len(resources)
        additional = general_resources[:remaining]
        resources.extend(additional)
    
    # Customize based on difficulty level
    for resource in resources:
        if difficulty.lower() == "beginner":
            if "Fundamentals" not in resource["title"] and "Basics" not in resource["title"]:
                resource["title"] = resource["title"] + " (Beginner-Friendly)"
        elif difficulty.lower() == "advanced":
            if "Advanced" not in resource["title"] and "Expert" not in resource["title"]:
                resource["title"] = "Advanced " + resource["title"]
    
    return resources


# Configure the API key for Google Generative AI
genai.configure(api_key="AIzaSyALGEs16JYUEYoeAjxHHKmgSODzVGEO8is")

# Function to generate question and answer
def generate_qa():
    try:
        # Initialize the Generative Model
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        # Prompt to generate question and answer pairs
        response = model.generate_content(
            "Generate 1 question of Artificial Intelligence on topic of Neural network or on Natural language processing with its answers. "
            "Provide the output in JSON format with keys 'question' and 'answer'.",
            
        )
        print(response.text)
        return response.text
    except Exception as e:
        return {"error": str(e)}


@app.route('/generate_path_recommendations', methods=['POST'])
def generate_path_recommendations():
    try:
        # Get data from request
        data = request.json
        
        if not data or 'skills' not in data:
            return jsonify({"error": "No skills data provided"}), 400
            
        skills = data.get('skills', [])
        
        # Configure Gemini API
        genai.configure(api_key="AIzaSyALGEs16JYUEYoeAjxHHKmgSODzVGEO8is")
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        # Create prompt based on skills
        skill_description = ", ".join([f"{skill['name']} (proficiency: {skill['proficiency']}%)" for skill in skills])
        
        prompt = f"""
        Based on the following skills and proficiency levels:
        {skill_description}
        
        Generate a personalized learning path recommendation with:
        1. Best career paths that match these skills (Frontend Developer, Backend Developer, Full Stack Developer, etc.)
        2. Learning resources (courses, books, tutorials) that would help improve these skills
        3. Additional skills to develop that would complement the existing ones
        
        Return the response in JSON format with three sections:
        - careerPaths: array of objects with title, description, skills (array), timeline, avgSalary, growthRate
        - learningResources: array of objects with title, type, provider, difficulty, url, rating
        - skillsToLearn: array of objects with name, priority ("High" or "Medium"), category ("Frontend" or "Backend")
        
        Make sure the JSON is valid and properly formatted.
        """
        
        # Generate content with Gemini
        response = model.generate_content(prompt)
        
        # Parse the response
        response_text = response.text
        
        # Clean the response to get properly formatted JSON
        if '```json' in response_text:
            response_text = response_text.split('```json')[1].split('```')[0].strip()
        elif '```' in response_text:
            response_text = response_text.split('```')[1].split('```')[0].strip()
            
        try:
            # Parse JSON response
            recommendations = json.loads(response_text)
            return jsonify(recommendations)
        except json.JSONDecodeError as e:
            # If parsing fails, return error with sample data
            print(f"JSON parse error: {str(e)}")
            print(f"Raw response: {response_text}")
            
            # Provide fallback data structure
            fallback_response = {
                "careerPaths": [
                    {
                        "title": "Frontend Developer",
                        "description": "Specialize in creating user interfaces and experiences for web applications.",
                        "skills": ["JavaScript", "React", "HTML/CSS", "TypeScript", "UI/UX Design"],
                        "timeline": "1-3 months",
                        "avgSalary": "$90,000 - $120,000",
                        "growthRate": "24% (Much faster than average)"
                    },
                    {
                        "title": "Full Stack Developer",
                        "description": "Develop both client and server software for complete web applications.",
                        "skills": ["JavaScript", "React", "Node.js", "SQL", "MongoDB", "AWS"],
                        "timeline": "4-6 months",
                        "avgSalary": "$100,000 - $130,000",
                        "growthRate": "22% (Much faster than average)"
                    }
                ],
                "learningResources": [
                    {
                        "title": "Advanced React Patterns",
                        "type": "Course",
                        "provider": "Frontend Masters",
                        "difficulty": "Intermediate",
                        "url": "https://frontendmasters.com/courses/advanced-react-patterns/",
                        "rating": 4.8
                    },
                    {
                        "title": "Node.js Design Patterns",
                        "type": "Book",
                        "provider": "O'Reilly",
                        "difficulty": "Advanced",
                        "url": "https://www.oreilly.com/library/view/nodejs-design-patterns/9781785885587/",
                        "rating": 4.6
                    }
                ],
                "skillsToLearn": [
                    {
                        "name": "TypeScript",
                        "priority": "High",
                        "category": "Frontend"
                    },
                    {
                        "name": "GraphQL",
                        "priority": "Medium",
                        "category": "Backend"
                    },
                    {
                        "name": "AWS",
                        "priority": "High",
                        "category": "Backend"
                    }
                ]
            }
            
            return jsonify(fallback_response)
        
    except Exception as e:
        print(f"Error generating path recommendations: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/start_interview', methods=['GET'])
def start_interview():
    print("Called")
    # Simulate running the interview process
    qa1 = generate_qa()
    
    # Clean and parse the response
    qa_cleaned = qa1[8:-4]
    try:
        qa = json.loads(qa_cleaned)
    except json.JSONDecodeError as e:
        return jsonify({"error": "Failed to parse the question."}), 500
    
    if "error" in qa:
        return jsonify({"error": qa["error"]}), 500
    global question,answer
    question = qa.get("question", "No question found.")
    answer = qa.get("answer","No answer found.")
    print("Answer returned: ",answer)
    print("Question returned: ",question)
    return jsonify({"question": question})



@app.route('/evaluate_answer', methods=['POST'])
def evaluate_answer():
    data = request.get_json()  
    global question,answer,user_answer
    selected = ""
    user_answer = data.get('answer', '')  
    print("Received Answer:", user_answer) 
    checkScore = calculate_similarity(answer,user_answer)
    if checkScore>0.3:
        selected = "Selected"
    else:
        selected = "Not Selected"
    data=[question,answer,user_answer,checkScore,selected]
    with open(file_name, mode="a", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(data)
    return jsonify({'status': 'success', 'message': 'Answer received and printed'})


# Run Flask app
if __name__ == "__main__":
    app.run(debug=True)
